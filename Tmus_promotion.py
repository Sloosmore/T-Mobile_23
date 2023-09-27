import pandas as pd
import re
from docx.shared import Pt
from docx import Document
from docx2pdf import convert
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.shared import Cm, Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import RGBColor
import os

import pandas as pd

# Assuming 'df' is your DataFrame
# Replace "-" with "." if it is surrounded by spaces on either side

###############################################################################
# This function calculates the number of words in given responsibility list
###############################################################################
def calc_nwords(resp):
    words = resp.split(" ")
    return len(words)

###############################################################################
# created 7/21 to deal with extra sentences and add those as separate responsibilities
###############################################################################
def process_resp(resplist, text):
    if text.count(".")>1:
        y = text.split(".")
#        print(y)
        for sent in y:
            sent = sent.strip()
#            print(sent)
            if sent != '' or sent != " " or sent !='\n' or sent != "." or sent != ". ":
                resplist.append(sent) 
    else:
        resplist.append(text)

    return resplist   

###############################################################################
# this function takes in a raw responsibility and attempts to figure out:
#  the title (if it exists) and the sub tasks.
###############################################################################
def process_responsibility(resp):
    # storage for the categoy and responsibility list
    cat = ""
    resplist = []
    # is there a dash with a space immediately after?  a dash with letters before AND after
    #  is probably just a hyphon so ignore that.  A dash with a space after is likely a delimiter
    if '-' in resp:
        # print("dashes found!")
        delim = []
        for i in range(0, len(resp)):
            # check if there is a space after, if so it's a delimiter
            if resp[i]=="-" and resp[i+1]==' ': 
                delim.append(i)
        # print out where the delimiters are just to check
        # print(delim)

        # now, if there are any delimiters, divide up the given string based on where 
        #  the delimiters occur.  Note if ther are 2 delimiters, there should be 3 resulting strings
        if (len(delim)>0):
            # add the overall string length
            delim.append(len(resp)+1)
            # check the first one, is it a category or a responsibility?
            first = resp[0:delim[0]]
            # somewhat arbitrary, but if the first has less than 5 words, assume it's actually
            #  the category
            if calc_nwords(first) <= 6: 
                cat = first
            # otherwise (if > 5 words) assume the first is a responsibility
            else:
                first = first.replace("- ", "")
                resplist.append(first)

            # now (regardless) go through the rest of the delimiters and append to resplist
            j = delim[0] 
            for k in delim[1:]:
                text = resp[j:k]
                text = text.replace("- ", "")
                # some processing on text before appending, e.g. to make sure it is not multiple sentences 
                resplist = process_resp(resplist, text)
                j=k

        else:
            # no category (but we should also check for colon, right?)
            # if no delimiters, then resplist has one element: the entirity of resp 
            resplist = process_resp(resplist, resp)
    else:
        # no dashes at all!  Again here resplist has one element: the entirity of resp 
        resplist = process_resp(resplist, resp)

    ## final check to make sure there aren't any blank entries ##
    for resp in resplist:
        if resp=='':
#            print("got here!")
            resplist.remove(resp)
    return cat, resplist

def generate_title(ex):
        # empty list that will contain the title
        title = []

        # start by splitting by the comma so we know front and back portions
        jobs = ex.split(", ")

        # if there was a comma then the length will be greater than 1
        if len(jobs) > 1:
            # store the front and back portions
            front = jobs[0]
            back = jobs[1]

            # Remove extra space before the first word in front portion, if present
            front = re.sub(r'^\s+', '', front).upper()

            # is either SR or PRINCIPAL in the front portion?
            if ("SR" in front) or ("PRINCIPAL" in front):
                # if so, reorder
                words = front.split()
                title.append(words[0].title())  # Apply title() method
                title.append(back.title())  # Apply title() method
                title.append(words[1].title())  # Apply title() method
            else:
                # otherwise, just switch the front and back portions
                title.append(back.title())  # Apply title() method
                title.append(front.title())  # Apply title() method

        # since title is now a list, make it a string instead
            title = ' '.join([str(elem) for elem in title])
            title = re.sub(r'^\s+', '', title)


        # if there wasn't a comma, then the title is just the input string
        else:
            title = ex

        return title

def set_font_and_size(paragraph, font_name, font_size):
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = font_size

def add_fillable_field(doc, label):
    p = doc.add_paragraph()
    p.add_run(label).bold = True
    set_font_and_size(p, "Calibri", Pt(12))

def apply_shading(paragraph):
    shading_elm = parse_xml(r'<w:shd {} w:fill="D3D3D3"/>'.format(nsdecls('w')))
    paragraph._element.get_or_add_pPr().append(shading_elm)



def csv_to_word(csv):
    df = pd.read_csv("all csv files/"+csv)
        # Select the desired row using iloc
    isolated_tasks = df.iloc[[2]].iloc[:, 2:]
    isolated_tasks.columns = range(len(isolated_tasks.columns))
    
    row_index = 0

    # Iterate over each cell in the row
    for col in isolated_tasks.columns:
        cell_value = isolated_tasks.iloc[row_index, col]
        if isinstance(cell_value, str):
            # Remove digits before %
            modified_value = re.sub(r'(\d+)(?=%)', '', cell_value)
            # Remove %
            modified_value = modified_value.replace('%', '\n')
            isolated_tasks.iloc[row_index, col] = modified_value
    #remove whitespaces before slice
    isolated_tasks = isolated_tasks.apply(lambda x: x.str.strip() if x.dtype == 'object' else x)
    isolated_tasks.replace('▪', '-', regex=True, inplace=True)
    isolated_tasks.replace('§', '', regex=True, inplace=True)
    isolated_tasks.replace(':', '', regex=True, inplace=True)
    isolated_tasks.replace('T-\nMobile', 'T-Mobile', regex=True, inplace=True)
    #isolated_tasks.replace('e.g.', 'for exsample', regex=False, inplace=True)
    #print(isolated_tasks)
    # Iterate over each cell in the row
    output_df = pd.DataFrame(columns=isolated_tasks.columns)
    # Iterate over each cell in the row
    for col in isolated_tasks.columns:
        cell_value = isolated_tasks.iloc[row_index, col]
        if isinstance(cell_value, str):
            # Split the text into sections based on a blank line
            sections = re.split(r'\n\s*\n', cell_value)
            print(sections)
            # Add each section as a new row in the output DataFrame
            for i, section in enumerate(sections):
                output_df.loc[i, col] = section.strip()
    print()
    #print(df)
    
    isolate_title = df.iloc[0, 1]

    title_list = isolate_title.split('\n')
    title_list = [title.title() for title in title_list]

    new_titles_df = pd.DataFrame([title_list]).drop(columns=0).reset_index(drop=True)
    new_titles_df.columns = range(len(new_titles_df.columns))
    
    fnl_titles_df = new_titles_df.applymap(generate_title)


    #combining both prior df's
    final_df = pd.concat([fnl_titles_df, output_df], ignore_index=True).reset_index(drop=True)
    #final_df.drop(final_df.index[-1], inplace=True)
    final_df = final_df.replace('\s{2,}', ' ', regex=True)
    
   
# Initialize the first_occurrence flag as True
    #final_df = final_df.applymap(global_replace)

    print(csv)
    print(final_df)

    for column in final_df.columns:
        # Create a new Word document
        doc = Document()
        
        doc.styles['Normal'].font.name = 'Calibri'
        doc.styles['Normal'].font.size = Pt(12)
        doc.styles['Normal'].paragraph_format.line_spacing = 1
        
        header = doc.sections[0].header
        htable=header.add_table(1, 1, Inches(6))
        htab_cells=htable.rows[0].cells
        ht0=htab_cells[0].add_paragraph()
        kh=ht0.add_run()
        kh.add_picture('Diglogo.png', width=Inches(6))
        ht0.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Get the column values
        column_values = final_df[column]

        # Set the heading as the title of the document
        title = str(column_values.iloc[0])

        # Set the heading as the title of the document
        heading = doc.add_heading(title + ' Promotion Packet', level=1)
        run = heading.runs[0]

    # Set the font color to black (RGB value: 0, 0, 0)
        run.font.color.rgb = RGBColor(0, 0, 0)

        # Set the font size and color of the title
        for paragraph in doc.paragraphs:
            run = paragraph.runs
            if run:
                run[0].font.size = Pt(16)  # Adjust the font size as needed
        
        doc.add_paragraph()
    
        add_fillable_field(doc, "Name:                                                                                           _________________________")
        add_fillable_field(doc, "Current Position:                                                                        _________________________")
        add_fillable_field(doc, "Email:                                                                                            _________________________")
        add_fillable_field(doc, "Manager:                                                                                      _________________________")
        add_fillable_field(doc, "DLT Member:                                                                               _________________________")
        
        you = doc.add_paragraph()
        you_run = you.add_run('Direct Manager Endorsement Comments:')
        you_run.bold = True
        apply_shading(you)
        
        question1 = doc.add_paragraph()
        run = question1.add_run("\nIC’s:")
        run.bold = True
        set_font_and_size(question1, "TeleNeo Office", Pt(12))
        question1.add_run(" This is your opportunity to showcase your accomplishments in your current role and readiness for the next. "
                            "A big part of being ready for the next role is the ability to self-assess on both your strengths and opportunities "
                            "and then showcase...\n" "Each section below is part of the core responsibilities from your desired role’s Job Matrix."
                                    " For each section describe in detail the ways that you are\n" 
                                    "   1)  Demonstrating these responsibilities in your current role\n"
                                    "   2)  Opportunity areas that you have identified and your action plans to grow into the role"
                            )
        

            # Add second question as a paragraph
            

            # Add spacing

            # Add third question as a paragraph
        question3 = doc.add_paragraph()
        run = question3.add_run("Direct Managers:")
        run.bold = True
        set_font_and_size(question3, "TeleNeo Office", Pt(12))
        question3.add_run(" Review this with your staff and provide your endorsement comments at the end.")

            # Add spacing
        doc.add_paragraph()
        # Iterate over the cells in the column (starting from index 1)
        z = 1
        for idx, value in enumerate(column_values[1:], start=1):
            # Add the value as a paragraph
            
            # Find the index of the first '-'
            value = str(value)
            cat, resplist = process_responsibility(value)
            if (len(cat)>0):
                # Get the words before the '-' and the '-' itself
                

                # Create a new paragraph
                paragraph = doc.add_paragraph()
                cat  = cat+':'
                cat = cat.replace(" :", ":")
                # Add the bold text
                run = paragraph.add_run(cat)
                run.bold = True

                # Add the remaining text
                resplist = [item.replace('TMobile', 'T-Mobile') for item in resplist]
                resplist = [item.strip() for item in resplist]
                resplist = [item.replace('.', '') for item in resplist]
                for x in resplist:
                    if not str(x) == '' or str(x) == 'nan':
                        #print(type(x))
                        x = str(x)
                        doc.add_paragraph(x+'.', style='ListBullet')


                apply_shading(paragraph)
                set_font_and_size(paragraph, "Calibri", Pt(12))

            else:
                # If there is no '-', add the entire value as a regular paragraph
                if not str(value) == 'nan':
                    paragraph = doc.add_paragraph()
                    run = paragraph.add_run('Main Responsibility '+ str(z))
                    run.bold = True
                    

                    apply_shading(paragraph)
                    resplist = [item.replace('TMobile', 'T-Mobile') for item in resplist]
                    resplist = [item.strip() for item in resplist]
                    resplist = [item.replace('.', '') for item in resplist]

                    for x in resplist:
                        if not str(x) == '' or str(x) == 'nan':
                            #print(type(x))
                            x = str(x)
                            doc.add_paragraph(x+'.', style='ListBullet')
                            apply_shading(paragraph)
                    set_font_and_size(paragraph, "Calibri", Pt(12))
            z += 1

            if not str(value) == 'nan':
                doc.add_paragraph()

                paragraph1 = doc.add_paragraph()
                
                run1 = paragraph1.add_run('A)   Describe in detail the ways that you are already operating at this level in your current role. '
                                        'Specific examples with impacts that are tied into a narrative are most helpful. '
                                        'Bring your abilities to life.')
                run1.bold = True
                
                
                

                # Create a table with a single cell
                table = doc.add_table(rows=1, cols=1)

                # Get the first (and only) cell in the table
                table.style = 'Table Grid'
                for row in table.rows:
                    row.height = Cm(6)

                doc.add_paragraph()

                # Add spacing

                # Second paragraph
                paragraph2 = doc.add_paragraph()
                run2 = paragraph2.add_run('B)   In what areas have you identified opportunities and what have you been doing to strengthen '
                                        'and develop while in role?')
                run2.bold = True
                
                table = doc.add_table(rows=1, cols=1)

                doc.add_paragraph()

                # Get the first (and only) cell in the table
                table.style = 'Table Grid'
                for row in table.rows:
                    row.height = Cm(6)
            else:
                break
            # Add spacing
            
        end = doc.add_paragraph()
        end_run = end.add_run('Direct Manager Endorsement Comments:')
        end_run.bold = True
        apply_shading(end)

        doc.add_paragraph('Please highlight in your words your assessment of the candidate readiness for desired role.')
        
        table = doc.add_table(rows=1, cols=1)
        table.style = 'Table Grid'
        for row in table.rows:
            row.height = Cm(6)
        

        doc.save(f'./Word output/{title} Promotion Packet.docx')

        # Convert the Word document to PDF
        #convert(f'{title} Promotion Packet.docx', f'{title} Promotion Packet.pdf')
        






# Assuming your DataFrame is named output_df
    first_cell = final_df.iloc[0, 0]  # Get the value of the first cell

    # Extract the first few words from the first cell
    first_few_words = " ".join(first_cell.split()[:3])  # Change '3' to the number of words you want in the title

    # Use the extracted words as the title for the document
    final_df.to_csv(f'./csv_check/{first_few_words} Promotion Packet.csv')

## initialize a blank list of files
csv_files = []

## loop through all the files and store the names of the pdf files
for x in os.listdir("all csv files"):
    if x.endswith(".csv"):
        csv_files.append(x)
#print(len(csv_files))
#print(csv_files[0])
for csv in csv_files:
    csv_to_word(csv)
